"""Word (.docx) extractor."""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document

from extractors.base import Extractor

logger = logging.getLogger(__name__)


class DocxExtractor(Extractor):
    """Extract raw text from .docx files.

    Reads paragraphs and tables, preserving reading order.
    """

    def extract(self, file_path: str | Path) -> str:
        doc = Document(str(file_path))
        parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append("\t".join(cells))

        return "\n".join(parts)
